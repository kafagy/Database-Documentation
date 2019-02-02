import pandas as pd
import mysql.connector
from docx import Document
from docx.shared import Mm

# Create Word Document Object
document = Document()
document.add_heading('Database Specification', 0)


conn = mysql.connector.connect(user='', password='',host='127.0.0.1',database='')
cursor = conn.cursor()

cursor.execute('''
SELECT TABLE_NAME 
FROM information_schema.tables t
WHERE t.table_schema = 'TRENDY'  -- put schema name here
ORDER BY t.table_name;
''')

tableNames = cursor.fetchall()
print(tableNames)
for tableName in tableNames:
    print(tableName[0])
    df = pd.read_sql_query('''
            SELECT ordinal_position AS position,
       column_name,
       data_type,
       CASE WHEN character_maximum_length IS NOT NULL
            then character_maximum_length
            else numeric_precision end AS max_length,
       is_nullable,
       column_default AS default_value
FROM information_schema.columns
WHERE table_name = '{}'
ORDER BY ordinal_position;
            '''.format(*tableName), conn)

    PK = pd.read_sql_query('''
            SELECT tab.table_schema AS database_schema,
    sta.index_name AS pk_name,
    sta.seq_in_index AS column_id,
    sta.column_name,
    tab.table_name
FROM information_schema.tables AS tab
INNER JOIN information_schema.statistics AS sta
        ON sta.table_schema = tab.table_schema
        AND sta.table_name = tab.table_name
        AND sta.index_name = 'primary'
WHERE tab.table_schema = 'schema' AND tab.table_name = '{}'
ORDER BY tab.table_name,
    column_id;
;'''.format(*tableName), conn)

    FK = pd.read_sql_query('''
            SELECT concat(fks.constraint_schema, '.', fks.table_name) AS foreign_table,
       '->' AS rel,
       concat(fks.unique_constraint_schema, '.', fks.referenced_table_name)
              AS primary_table,
       fks.constraint_name,
       group_concat(kcu.column_name
            ORDER BY position_in_unique_constraint separator ', ') AS fk_columns
FROM information_schema.referential_constraints fks
join information_schema.key_column_usage kcu
     ON fks.constraint_schema = kcu.table_schema
     AND fks.table_name = kcu.table_name
     AND fks.constraint_name = kcu.constraint_name
WHERE fks.table_name = '{}' AND fks.constraint_schema = 'schema'
GROUP BY fks.constraint_schema,
         fks.table_name,
         fks.unique_constraint_schema,
         fks.referenced_table_name,
         fks.constraint_name
ORDER BY fks.constraint_schema,
         fks.table_name;
'''.format(*tableName), conn)

    # Max Length Of Every Column For Every DataFrame
    tableMax = [df[col].astype(str).apply(len).max() for col in df.columns]
    pkMax = [PK[col].astype(str).apply(len).max() for col in PK.columns]
    fkMax = [FK[col].astype(str).apply(len).max() for col in FK.columns]

    print(str(df) + '\n\n' + str(PK) + '\n\n' + str(FK))

    # Constructing Word Table For Every DataFrame
    document.add_heading('schema.' + tableName[0].upper(), level=1)
    t = document.add_table(df.shape[0] + 1, df.shape[1], style='Light Grid Accent 1')
    for idx in range(len(df.columns)):
        t.columns[idx].width = Mm(tableMax[idx] * 5)
        for cell in t.column_cells(idx):
            cell.width = Mm(tableMax[idx] * 5)
    for j in range(df.shape[-1]):
        t.cell(0, j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i + 1, j).text = str(df.values[i, j])

    if not PK.empty:
        document.add_heading('schema.' + tableName[0].upper() + ' - Primary Keys', level=1)
        t = document.add_table(PK.shape[0] + 1, PK.shape[1], style='Light Grid Accent 1')
        for idx in range(len(PK.columns)):
            t.columns[idx].width = Mm(pkMax[idx] * 3)
            for cell in t.column_cells(idx):
                cell.width = Mm(pkMax[idx] * 3)
        for j in range(PK.shape[-1]):
            t.cell(0, j).text = PK.columns[j]
        for i in range(PK.shape[0]):
            for j in range(PK.shape[-1]):
                t.cell(i + 1, j).text = str(PK.values[i, j])

    if not FK.empty:
        document.add_heading('schema.' + tableName[0].upper()  + ' - Foreign Keys', level=1)
        t = document.add_table(FK.shape[0] + 1, FK.shape[1], style='Light Grid Accent 1')
        for idx in range(len(FK.columns)):
            t.columns[idx].width = Mm(fkMax[idx] * 3)
            for cell in t.column_cells(idx):
                cell.width = Mm(fkMax[idx] * 3)
        for j in range(FK.shape[-1]):
            t.cell(0, j).text = FK.columns[j]
        for i in range(FK.shape[0]):
            for j in range(FK.shape[-1]):
                t.cell(i + 1, j).text = str(FK.values[i, j])

# Saving Word Document
document.save('MySQL_Specs.docx')
