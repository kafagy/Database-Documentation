import pyodbc
import pandas as pd
from docx import Document
from docx.shared import Mm

# Create Word Document Object
document = Document()
document.add_heading('Database Specification', 0)

# Set Printing Configurations For Pandas
pd.set_option('display.width', 1000)
pd.set_option('display.height', 1000)
pd.set_option('display.max_rows', 1000)
pd.set_option('display.max_columns', 1000)

# Database Connection
conn = pyodbc.connect(
            r'DRIVER={ODBC Driver 13 for SQL Server};'
            r'SERVER=;'
            r'DATABASE=;'
            r'UID=;'
            r'PWD='
)
cursor = conn.cursor()
cursor.execute('''
          SELECT 
            T.TABLE_NAME AS 'Table'
          FROM INFORMATION_SCHEMA.TABLES T
          WHERE TABLE_TYPE = 'BASE TABLE' 
            AND TABLE_SCHEMA = 'schema'
          ORDER BY T.TABLE_NAME ASC;''')
tableNames = cursor.fetchall()
for tableName in tableNames:
    df = pd.read_sql_query('''
            SELECT
               C.ORDINAL_POSITION AS 'No', 
               C.COLUMN_NAME AS 'Column',
               CASE C.DATA_TYPE
                WHEN 'varchar' THEN CONCAT('Varchar(', C.CHARACTER_MAXIMUM_LENGTH, ')')
                WHEN 'char'    THEN CONCAT('Char(', C.CHARACTER_MAXIMUM_LENGTH, ')')
                WHEN 'decimal' THEN CONCAT('Decimal(', C.NUMERIC_PRECISION, ', ', C.NUMERIC_SCALE, ')')
                WHEN 'int'     THEN 'Integer(10)'
                ELSE C.DATA_TYPE
               END AS 'Datatype',
               C.COLUMN_DEFAULT AS 'Default',
               C.IS_NULLABLE AS 'Null'
            FROM INFORMATION_SCHEMA.TABLES T
            INNER JOIN INFORMATION_SCHEMA.COLUMNS C
              ON C.TABLE_NAME = T.TABLE_NAME
              AND T.TABLE_SCHEMA = C.TABLE_SCHEMA
            WHERE T.TABLE_TYPE='BASE TABLE'
              AND T.TABLE_SCHEMA = 'schema'
              AND T.TABLE_NAME = '{}'
            ORDER BY T.TABLE_NAME;'''.format(*tableName), conn)

    PK = pd.read_sql_query('''
            SELECT DISTINCT
                  KU.CONSTRAINT_NAME AS 'Constraint',
                  KU.COLUMN_NAME AS 'PK Column'
            FROM INFORMATION_SCHEMA.TABLE_CONSTRAINTS TC
            INNER JOIN INFORMATION_SCHEMA.KEY_COLUMN_USAGE KU
                  ON TC.CONSTRAINT_TYPE = 'PRIMARY KEY'
                  AND TC.CONSTRAINT_NAME = KU.CONSTRAINT_NAME
                  AND KU.TABLE_NAME='{}';'''.format(*tableName), conn)

    FK = pd.read_sql_query('''
            SELECT DISTINCT 
                C.CONSTRAINT_NAME AS 'Constraint',
                FK.TABLE_NAME AS 'FK Table',
                CU.COLUMN_NAME AS 'FK Column',
                PK.TABLE_NAME AS 'PK Table',
                PT.COLUMN_NAME AS 'PK Column'
            FROM INFORMATION_SCHEMA.REFERENTIAL_CONSTRAINTS C
            INNER JOIN INFORMATION_SCHEMA.TABLE_CONSTRAINTS FK
                  ON C.CONSTRAINT_NAME = FK.CONSTRAINT_NAME
            INNER JOIN INFORMATION_SCHEMA.TABLE_CONSTRAINTS PK
                  ON C.UNIQUE_CONSTRAINT_NAME = PK.CONSTRAINT_NAME
            INNER JOIN INFORMATION_SCHEMA.KEY_COLUMN_USAGE CU
                  ON C.CONSTRAINT_NAME = CU.CONSTRAINT_NAME
            INNER JOIN (
                        SELECT
                            INNER1.TABLE_NAME,
                            INNER2.COLUMN_NAME
                        FROM INFORMATION_SCHEMA.TABLE_CONSTRAINTS INNER1
                        INNER JOIN INFORMATION_SCHEMA.KEY_COLUMN_USAGE INNER2
                              ON INNER1.CONSTRAINT_NAME = INNER2.CONSTRAINT_NAME
                      WHERE INNER1.CONSTRAINT_TYPE = 'PRIMARY KEY'
                 ) PT
                 ON PT.TABLE_NAME = PK.TABLE_NAME
          WHERE FK.TABLE_NAME = '{}'
            AND FK.TABLE_SCHEMA = 'schema'
            AND PK.TABLE_SCHEMA = 'schema'
            AND CU.TABLE_SCHEMA = 'schema';'''.format(*tableName), conn)

    # Pandas Manipulation
    df['Default'] = df['Default'].replace('[()]', '', regex=True)

    # Max Length Of Every Column For Every DataFrame
    tableMax = [df[col].astype(str).apply(len).max() for col in df.columns]
    pkMax = [PK[col].astype(str).apply(len).max() for col in PK.columns]
    fkMax = [FK[col].astype(str).apply(len).max() for col in FK.columns]

    # DataFrame Printing Statement
    print(str(df) + '\n\n' + str(PK) + '\n\n' + str(FK))

    # Constructing Word Table For Every DataFrame
    document.add_heading('schema.' + tableName[0], level=1)
    t = document.add_table(df.shape[0]+1, df.shape[1], style='Light Grid Accent 1')
    for idx in range(len(df.columns)):
        t.columns[idx].width = Mm(tableMax[idx] * 5)
        for cell in t.column_cells(idx):
            cell.width = Mm(tableMax[idx] * 5)
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    if not PK.empty:
        document.add_heading('schema.' + tableName[0] + ' - Primary Keys', level=1)
        t = document.add_table(PK.shape[0]+1, PK.shape[1], style='Light Grid Accent 1')
        for idx in range(len(PK.columns)):
            t.columns[idx].width = Mm(pkMax[idx] * 3)
            for cell in t.column_cells(idx):
                cell.width = Mm(pkMax[idx] * 3)
        for j in range(PK.shape[-1]):
            t.cell(0,j).text = PK.columns[j]
        for i in range(PK.shape[0]):
            for j in range(PK.shape[-1]):
                t.cell(i+1,j).text = str(PK.values[i,j])
    if not FK.empty:
        document.add_heading('schema.' + tableName[0]  + ' - Foreign Keys', level=1)
        t = document.add_table(FK.shape[0]+1, FK.shape[1], style='Light Grid Accent 1')
        for idx in range(len(FK.columns)):
            t.columns[idx].width = Mm(fkMax[idx] * 3)
            for cell in t.column_cells(idx):
                cell.width = Mm(fkMax[idx] * 3)
        for j in range(FK.shape[-1]):
            t.cell(0,j).text = FK.columns[j]
        for i in range(FK.shape[0]):
            for j in range(FK.shape[-1]):
                t.cell(i+1,j).text = str(FK.values[i,j])

# Saving Word Document
document.save('SQLSERVER_Specs.docx')
