import ibm_db_dbi
import pandas as pd
from docx import Document
from docx.shared import Mm

# Create Word Document Object
document = Document()
document.add_heading('Database Specification', 0)

# Set Printing Configurations For Pandas
pd.set_option('display.width', 1000)
pd.set_option('display.height', 1000)
pd.set_option('display.max_rows', 1000)
pd.set_option('display.max_columns', 1000)

conn = ibm_db_dbi.connect("DATABASE=;HOSTNAME=;PORT=;PROTOCOL=;UID=;PWD=;", "", "")
cursor = conn.cursor()
cursor.execute('''
            SELECT
              T.TABNAME AS Table
            FROM
              SYSCAT.TABLES T
            WHERE T.TABSCHEMA = ''
              AND T.TYPE = 'T'
            ORDER BY T.TABNAME;''')
tableNames = cursor.fetchall()
for tableName in tableNames:
    print(tableName[0])
    df = pd.read_sql_query('''
            SELECT
              C.COLNO + 1 AS No,
              C.COLNAME AS Column,
              CASE C.TYPENAME
                WHEN 'VARCHAR' THEN 'Varchar(' || C.LENGTH || ')'
                WHEN 'CHAR'    THEN 'Char(' || C.LENGTH || ')'
                WHEN 'DECIMAL' THEN 'Decimal(' || C.LENGTH || ', ' || C.SCALE || ')'
                WHEN 'INTEGER' THEN 'Integer(10)'
                WHEN 'TIME' THEN 'Time'
                WHEN 'BIGINT' THEN 'BigInt'
                WHEN 'DATE' THEN 'Date'
              ELSE C.TYPENAME
              END AS Datatype,
              C.DEFAULT AS Default,
              C.NULLS AS Null
            FROM
              SYSCAT.TABLES T
            INNER JOIN SYSCAT.COLUMNS C
              ON T.TABNAME = C.TABNAME
              AND T.TABSCHEMA = C.TABSCHEMA
            WHERE C.TABSCHEMA = 'schema'
              AND T.TABNAME = '{}'
              AND TYPE = 'T'
            ORDER BY No;
            '''.format(*tableName), conn)

    PK = pd.read_sql_query('''
            SELECT
              SI.Name AS Constraint,
              REPLACE(LTRIM(SI.COLNAMES, '+'), '+', ' ') AS Column,
              SI.UNIQUERULE AS "Unique Rule"
            FROM
              SYSIBM.SYSINDEXES SI
            WHERE SI.TBNAME = '{}'
              AND SI.TBCREATOR = 'schema';'''.format(*tableName), conn)

    FK = pd.read_sql_query('''
            SELECT
              R.CONSTNAME AS Constraint,
              R.TABNAME AS "FK Table",
              R.FK_COLNAMES AS "FK Column",
              R.REFTABNAME AS "PK Table",
              R.PK_COLNAMES AS "PK Column"
            FROM
              SYSCAT.REFERENCES R
            WHERE TABNAME = '{}'
              AND TABSCHEMA = 'schema';'''.format(*tableName), conn)

    # Max Length Of Every Column For Every DataFrame
    tableMax = [df[col].astype(str).apply(len).max() for col in df.columns]
    pkMax = [PK[col].astype(str).apply(len).max() for col in PK.columns]
    fkMax = [FK[col].astype(str).apply(len).max() for col in FK.columns]

    print(str(df) + '\n\n' + str(PK) + '\n\n' + str(FK))

    # Constructing Word Table For Every DataFrame
    document.add_heading('schema.' + tableName[0], level=1)
    t = document.add_table(df.shape[0]+1, df.shape[1], style='Light Grid Accent 1')
    for idx in range(len(df.columns)):
        t.columns[idx].width = Mm(tableMax[idx] * 5)
        for cell in t.column_cells(idx):
            cell.width = Mm(tableMax[idx] * 5)
    for j in range(df.shape[-1]):
        t.cell(0,j).text = df.columns[j]
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            t.cell(i+1,j).text = str(df.values[i,j])
    if not PK.empty:
        document.add_heading('schema.' + tableName[0] + ' - Primary Keys', level=1)
        t = document.add_table(PK.shape[0]+1, PK.shape[1], style='Light Grid Accent 1')
        for idx in range(len(PK.columns)):
            t.columns[idx].width = Mm(pkMax[idx] * 3)
            for cell in t.column_cells(idx):
                cell.width = Mm(pkMax[idx] * 3)
        for j in range(PK.shape[-1]):
            t.cell(0,j).text = PK.columns[j]
        for i in range(PK.shape[0]):
            for j in range(PK.shape[-1]):
                t.cell(i+1,j).text = str(PK.values[i,j])
    if not FK.empty:
        document.add_heading('schema.' + tableName[0]  + ' - Foreign Keys', level=1)
        t = document.add_table(FK.shape[0]+1, FK.shape[1], style='Light Grid Accent 1')
        for idx in range(len(FK.columns)):
            t.columns[idx].width = Mm(fkMax[idx] * 3)
            for cell in t.column_cells(idx):
                cell.width = Mm(fkMax[idx] * 3)
        for j in range(FK.shape[-1]):
            t.cell(0,j).text = FK.columns[j]
        for i in range(FK.shape[0]):
            for j in range(FK.shape[-1]):
                t.cell(i+1,j).text = str(FK.values[i,j])

# Saving Word Document
document.save('DB2_Specs.docx')
